import sys
import json
import docx

from argparse import ArgumentParser


def argparser():
    ap = ArgumentParser()
    ap.add_argument('jsonl', help='original-data/databricks-dolly-15k.jsonl')
    ap.add_argument('--max-len', type=int, default=950000)
    return ap


def yield_data(fname):
    with open(fname) as inp:
        for line in inp:
            yield json.loads(line)


def main(argv):
    args = argparser().parse_args(argv[1:])

    doc = docx.Document()
    doc_counter = 0
    current_len = 0

    for d_idx, d in enumerate(yield_data(args.jsonl)):
        d_len=len(d["instruction"])+len(d["context"])+len(d["response"])+50
        if (current_len+d_len) > args.max_len:
            doc.save(f"dolly-doc-in/dolly-{doc_counter:05d}.docx")
            doc_counter += 1
            current_len = 0
            doc = docx.Document()

        pgraph = doc.add_paragraph("")

        r = pgraph.add_run(f"Document {d_idx}")
        r.bold = True
        r.underline = True

        pgraph=doc.add_paragraph("")
        r = pgraph.add_run(f"Instruction")
        r.bold=True
        pgraph = doc.add_paragraph(d["instruction"])

        if d["context"] and not d["context"].isspace():
            pgraph=doc.add_paragraph("")
            r = pgraph.add_run(f"Context")
            r.bold=True
            pgraph = doc.add_paragraph(d["context"])

        pgraph=doc.add_paragraph("")
        r = pgraph.add_run(f"Response")
        r.bold=True
        pgraph = doc.add_paragraph(d["response"])

        current_len+=d_len
    else:
        doc.save(f"dolly-doc-in/dolly-{doc_counter:05d}.docx")


if __name__=="__main__":
    sys.exit(main(sys.argv))
